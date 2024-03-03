from opencage.geocoder import OpenCageGeocode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Cm
import src.utils as utils
import requests
import pandas as pd


class cordinates:
    def __init__(self):
        pass


    def get_location_by_address(address):
        geocoder = OpenCageGeocode('b7ce06f8bf404ce2957a8260aa9cf49e')
        results = geocoder.geocode(address)
        if results and len(results):
            print(results[0]['components']['city'])
            return (results[0]['geometry']['lat'], results[0]['geometry']['lng'], results[0]['components']['city'])
        
    def makedoc(data):
        def add_table(doc, table_data, column_headers):
            table = doc.add_table(rows=1, cols=len(column_headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(column_headers):
                hdr_cells[i].text = header
            for row_data in table_data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data):
                    row_cells[i].text = str(value)
                    row_cells[i].width = Cm(4)  # Set the width of each cell to 4 centimeters

        def add_paragraph(doc, text):
            paragraph = doc.add_paragraph(text)

        def generate_report(data):
            doc = Document()

            title = (data["title"])
            content = (data["content"])
            from_id = (data["from_id"])
            to_id = (data["to_id"])
            date = (data["date"])
            money = (data["money"])
            from_name = (data["from_name"])
            to_name = (data["to_name"])

            # Add title with centered alignment
            title_paragraph = doc.add_paragraph(title)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Make title light blue
            title_run = title_paragraph.runs[0]
            title_run.bold = True
            title_run.font.size = Pt(20)  # Adjust font size for the title
            title_run.font.color.rgb = RGBColor(135, 206, 250)  # Light blue color

            # Add paragraph spacing
            doc.add_paragraph().add_run().add_break()  # Add 2 line breaks

            # Create a table with the specified data
            table_data = list(zip(date, from_name, from_id, to_name,to_id, money))
            column_headers = ['Date', "Name", 'From', "Name",'To', 'Money']
            add_table(doc, table_data, column_headers)

            # Add paragraph spacing
            doc.add_paragraph().add_run().add_break()  # Add 2 line breaks

            # Add content as a paragraph
            add_paragraph(doc, content)

            doc.save(utils.REPORT) 

        # Example data
        generate_report(data)

        return utils.REPORT

    def callu(number):
        headers = {
        'Authorization': utils.NO
        }

        # Data
        data = {
        'phone_number': "+91"+ number,
        'task': 'real estate broker',
        'voice_id': 1,
        'reduce_latency': True,
        'request_data': {},
        'voice_settings':{
            "speed": "0.8"
        },
        'interruption_threshold': 0,
        'start_time': None,
        'transfer_phone_number': None,
        'answered_by_enabled': False,
        'from': None,
        'first_sentence': None,
        'record': True,
        'max_duration': 2,
        'model': 'enhanced',
        'language': 'ENG',
        }

        # API request 
        requests.post('https://api.bland.ai/call', json=data, headers=headers)   
        
        return "0"
        
    def reg_user(data):
            df = pd.read_csv(utils.DATABASE) 
            next_index = df.index.max() + 1 if not df.empty else 0
            
            df.loc[next_index, "wallet_address"] = data["wallet_address"]

            df.to_csv(utils.DATABASE, index=False)  
            
    def login_user(data):
        check = str(data["wallet_address"])
        # data = {"wallet_address": "12"}
        check = str(data["wallet_address"])
        df = pd.read_csv(utils.DATABASE)

        if df[df['wallet_address'] == check].empty:
            return "0"
        else:
            return "1"
            
    def returnid(data):
        print(data)
        city = data["city"]
        df = pd.read_csv(utils.LOCATION)
        df1 = df[df['city'] == city]  # Filter the DataFrame based on the city name

        id = df1["id"].values.tolist()
        add = df1["address"].values.tolist()
        lat = df1["lat"].values.tolist()
        long = df1["long"].values.tolist()
        city1 = df1["city"].values.tolist()

        data1 = {
            "id": id,
            "address": add,
            "lat": lat,
            "long": long,
            "city": city1,
        }

        return data1

        
    def enterid(data):
        def get_location_by_address(address):
            geocoder = OpenCageGeocode('b7ce06f8bf404ce2957a8260aa9cf49e')
            results = geocoder.geocode(address)
            if results and len(results):
                print(results[0]['components']['city'])
                return (results[0]['geometry']['lat'], results[0]['geometry']['lng'], results[0]['components']['city'])
        x, y, city = get_location_by_address(data["address"])
        print(x, y, city)
        
        data = {
            "id": data["id"],
            "address": data["address"],
            "lat": x,
            "long": y,
            "city": city,
        }

        df = pd.read_csv(utils.LOCATION)

        df.loc[len(df)] = data
    
        df.to_csv(utils.LOCATION, index=False)
        
        return "0"

